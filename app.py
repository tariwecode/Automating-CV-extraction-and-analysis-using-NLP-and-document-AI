import os
import re
import io
import math
import time
import sqlite3
import warnings
import shutil
from pathlib import Path
from datetime import datetime

import streamlit as st
import pandas as pd
import fitz  # PyMuPDF
import filetype
from docx import Document as DocxDocument
from dotenv import load_dotenv

import google.generativeai as genai
from google.api_core.exceptions import ResourceExhausted, NotFound

import chromadb
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.embeddings.fastembed import FastEmbedEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document as LcDocument
from pypdf.errors import PdfStreamError

from dashboard import render_dashboard
from auth import (
    authenticate_user, create_user, generate_reset_token, 
    reset_password, validate_email, validate_password, validate_username
)

@st.cache_data(show_spinner=False, max_entries=200)
def cached_extract_text(file_bytes, file_type):
    return extract_text(file_bytes, file_type)

@st.cache_data(show_spinner=False, max_entries=200)
def cached_embed(text):
    return embedding_model.embed_documents([text[:1500]])[0]

@st.cache_data(show_spinner=False, max_entries=200)
def cached_chunks(file_path):
    return extract_CV_chunks_from_path(file_path)

@st.cache_data(show_spinner=False, max_entries=200)
def cached_faiss(filename, chunks):
    return robust_create_or_load_vector_store(filename, chunks)

@st.cache_data(show_spinner=False, max_entries=200)
def cached_retrieve(job_desc, vectorstore):
    return retrieve_matching_chunks(job_desc, vectorstore)


st.set_page_config(page_title="CV Analysis System", page_icon="📄", layout="wide")

# Initialize authentication session state
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
    st.session_state.user = None
    st.session_state.auth_page = "login"  # login, signup, reset_password

load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
warnings.filterwarnings("ignore")

MODEL_NAME = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")
def _get_gemini_model():
    try:
        return genai.GenerativeModel(MODEL_NAME)
    except Exception:
        return None

GEMINI_MODEL = _get_gemini_model()

VECTOR_DIR = "chroma_store2"
UPLOAD_DIR = "temp_files1"
DB_PATH = "CV_data1.db"
embedding_model = FastEmbedEmbeddings(model_name="BAAI/bge-small-en-v1.5")

os.makedirs(VECTOR_DIR, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)

if "initialized" not in st.session_state:
    st.session_state.initialized = True
    st.session_state.submitted = False
    st.session_state.job_desc = ""
    st.session_state.uploaded_files = []
    st.session_state.CV_texts = {}
    st.session_state.analysis_results = []
    st.session_state.match_percentages = {}

conn = sqlite3.connect(DB_PATH)
c = conn.cursor()
c.execute('''CREATE TABLE IF NOT EXISTS CVs (
    filename TEXT PRIMARY KEY,
    candidate_name TEXT,
    match_percent TEXT,
    word_count INTEGER,
    skills TEXT,
    match_explanation TEXT
)''')
conn.commit()


def clear_cache(flag=1):
    c.execute("DELETE FROM CVs")
    conn.commit()

    if os.path.exists(VECTOR_DIR) and os.listdir(VECTOR_DIR):
        try:
            chroma_client = chromadb.PersistentClient(path=VECTOR_DIR)
            chroma_client.reset()
        except Exception:
            pass
        shutil.rmtree(VECTOR_DIR, ignore_errors=True)

    shutil.rmtree(UPLOAD_DIR, ignore_errors=True)
    os.makedirs(VECTOR_DIR, exist_ok=True)
    os.makedirs(UPLOAD_DIR, exist_ok=True)

    st.session_state.submitted = False
    st.session_state.job_desc = ""
    st.session_state.uploaded_files = []
    st.session_state.CV_texts = {}
    st.session_state.analysis_results = []
    st.session_state.match_percentages = {}
    st.session_state.jd_skills = []

    st.session_state.uploader_key = st.session_state.get("uploader_key", 0) + 1
    st.session_state.jd_key = st.session_state.get("jd_key", 0) + 1

    if not flag:
        st.success("Cleared successfully.")

    st.rerun()


def save_uploaded_file(file):
    base, ext = os.path.splitext(file.name)
    file_path = os.path.join(UPLOAD_DIR, file.name)
    counter = 1
    while os.path.exists(file_path):
        file_path = os.path.join(UPLOAD_DIR, f"{base}_{counter}{ext}")
        counter += 1
    with open(file_path, "wb") as f:
        f.write(file.getbuffer())
    return file_path

def extract_CV_chunks_from_path(file_path):
    try:
        loader = PyPDFLoader(file_path)
        documents = loader.load()
    except PdfStreamError:
        pdf = fitz.open(file_path)
        full_text = "\n".join(page.get_text() for page in pdf)
        documents = [LcDocument(page_content=full_text)]

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=2000,
        chunk_overlap=100,
        separators=["\n\n", "\n", ".", " "],
    )
    return splitter.split_documents(documents)

def retrieve_matching_chunks(job_description, vectorstore, k=3):
    return vectorstore.as_retriever(search_kwargs={"k": k}).get_relevant_documents(job_description)


def get_gemini_response(prompt, retrieved_chunks, job_description):
    context = "\n".join([chunk.page_content for chunk in (retrieved_chunks or [])])
    model = GEMINI_MODEL
    if model is None:
        st.error(f"Gemini model '{MODEL_NAME}' is unavailable. Set GEMINI_MODEL to a supported model from genai.list_models().")
        return ""  # no local fallback

    try:
        response = model.generate_content([
            prompt,
            f"Relevant CV Info:\n{context}",
            f"Job Description:\n{job_description}"
        ])
        return (response.text or "").strip()
    except NotFound:
        st.error(f"Model '{MODEL_NAME}' not found for this API. Set GEMINI_MODEL to a supported value (e.g., 'gemini-1.5-flash-latest' or 'gemini-1.5-pro-latest').")
        return ""
    except ResourceExhausted:
        st.error("Gemini quota exhausted. Try again later or upgrade your plan.")
        return ""
    except Exception as e:
        st.error(f"Gemini error: {e}")
        return ""


def extract_text(file_bytes, file_type):
    if file_type == "application/pdf" or file_bytes[:4] == b"%PDF":
        text_pages = []
        with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
            for page in pdf:
                text_pages.append(page.get_text())
        return "\n".join(text_pages)

    if "wordprocessingml.document" in file_type or file_type == "application/msword":
        doc = DocxDocument(io.BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs)

    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="ignore")


def count_matching_skills(text, skills):
    text_lower = text.lower()
    matched_skills = [skill for skill in skills if skill.lower() in text_lower]
    return len(matched_skills), matched_skills


def extract_skills_from_cv(cv_text: str) -> list:
    """Extract all identifiable skills from CV text."""
    if not cv_text:
        return []
    
    text_lower = cv_text.lower()
    found_skills = set()
    
    all_skills = [
        "python", "java", "javascript", "typescript", "c++", "c#", "ruby", "php", "swift", "kotlin",
        "sql", "mysql", "postgresql", "mongodb", "oracle", "nosql", "redis",
        "excel", "word", "powerpoint", "microsoft office", "ms office", "outlook",
        "html", "css", "react", "angular", "vue", "node.js", "django", "flask",
        "aws", "azure", "gcp", "docker", "kubernetes", "linux", "git", "github",
        "data analysis", "data entry", "data visualization", "machine learning",
        "accounting", "bookkeeping", "budgeting", "financial analysis", "auditing",
        "project management", "agile", "scrum", "jira", "trello",
        "marketing", "digital marketing", "seo", "social media", "content creation",
        "sales", "customer service", "client relations", "crm", "salesforce",
        "administration", "filing", "typing", "scheduling", "calendar management",
        "communication", "writing", "editing", "research", "reporting", "documentation",
        "planning", "organizing", "coordination", "event planning",
        "leadership", "teamwork", "team management", "mentoring", "coaching",
        "problem solving", "critical thinking", "analytical skills", "decision making",
        "time management", "attention to detail", "multitasking", "prioritization",
        "flexibility", "adaptability", "creativity", "innovation", "initiative",
        "interpersonal", "negotiation", "presentation", "public speaking",
        "training", "teaching", "tutoring", "facilitation",
        "graphic design", "photoshop", "illustrator", "figma", "canva", "ui/ux",
        "video editing", "photography", "content writing", "copywriting",
        "human resources", "hr", "recruitment", "talent acquisition", "onboarding",
        "logistics", "supply chain", "inventory management", "procurement",
        "quality assurance", "qa", "testing", "compliance", "risk management",
    ]
    
    for skill in all_skills:
        if skill in text_lower:
            found_skills.add(skill.title() if len(skill) > 3 else skill.upper())
    
    return sorted(list(found_skills))[:15]


prompt_analysis = """
You are a ATS (Applicant Tracking System). Your task is to analyze the following CV against the provided job description and highlight the key aspects in a crisp and to-the-point manner.

Highlight candidate Name.

Your analysis should be structured into the following four sections:

1.  **Profile Summary:** Provide a concise (1-2 sentences) summary of the candidate's core experience and key skills relevant to the job description.
2.  **Strengths:** List key strengths of the candidate both technical and soft skills in bullet points, focusing on direct relevance to the job description's requirements and preferences.
3.  **Weaknesses:** Identify potential weaknesses or areas for development in bullet points. Frame these objectively and professionally, focusing on gaps compared to the job description.
4.  **Missing Important Skills:** List critical skills explicitly mentioned in the job description that are not evident in the CV in bullet points. Be specific about the missing skills.

Ensure your analysis is concise, apt, and directly related to the information provided in the job description and CV. Avoid unnecessary elaboration or subjective opinions.
"""

prompt_match = """
You are an ATS (Applicant Tracking System). Analyze the following CV against the provided job description and provide a structured output.

Highlight candidate Name.

Your output should include the following three sections:

1.  **Matching Percentage:** Calculate and return a numerical percentage representing the overall match between the CV and the job description based on keyword presence and relevance.
2.  **Missing or Unmatched Keywords:** List the keywords and key phrases explicitly mentioned in the job description that are either absent or not prominently featured (unmatched) in the CV.
3.  **Final Thoughts:** Provide a brief (1-2 sentences) overall assessment of the candidate's suitability based on the keyword analysis.
"""


def extract_required_skills(job_description):
    """Extract skills from job description using Gemini, with local fallback."""
    try:
        model = GEMINI_MODEL
        if not model:
            raise RuntimeError("Model unavailable")

        prompt = (
            "Extract all key skills, qualifications, and competencies mentioned in the job description. "
            "Include soft skills. Return only a comma-separated list (no extra words).\n\n"
            f"{job_description}"
        )
        resp = model.generate_content(prompt)
        if resp and resp.text:
            skills = [s.strip().lower() for s in resp.text.split(",") if s.strip()]
            if skills:
                return skills
    except Exception:
        pass
    
    return _extract_skills_local(job_description)


def _extract_skills_local(text: str) -> list:
    """Local skill extraction using pattern matching."""
    if not text:
        return []
    
    text_lower = text.lower()
    found_skills = []
    
    common_skills = [
        "python", "java", "javascript", "sql", "excel", "word", "powerpoint",
        "data analysis", "data entry", "accounting", "bookkeeping", "budgeting",
        "project management", "marketing", "sales", "customer service",
        "administration", "filing", "typing", "microsoft office", "ms office",
        "communication", "writing", "editing", "research", "reporting",
        "scheduling", "planning", "organizing", "coordination",
        "leadership", "teamwork", "problem solving", "critical thinking",
        "time management", "attention to detail", "multitasking", "flexibility",
        "adaptability", "creativity", "initiative", "self-motivated",
        "interpersonal", "negotiation", "presentation", "public speaking",
    ]
    
    for skill in common_skills:
        if skill in text_lower:
            found_skills.append(skill)
    
    words = re.findall(r'\b[A-Z][a-zA-Z]+\b', text)
    for word in words:
        w_lower = word.lower()
        if len(w_lower) > 2 and w_lower not in _STOPWORDS and w_lower not in found_skills:
            found_skills.append(w_lower)
    
    return list(set(found_skills))[:30]


_STOPWORDS = {
    "and","or","the","a","an","to","of","in","with","for","on","at","by","as","be","is","are",
    "you","we","our","your","will","experience","years","year","minimum","good","great",
    "team","skills","requirements","responsibilities","qualifications","preferred","must",
    "should","who","what","when","where","why","how"
}

_SOFT_SKILLS = {
    "communication", "leadership", "teamwork", "adaptability", "empathy", "problem solving",
    "time management", "creativity", "attention to detail", "reliability", "customer service",
    "conflict resolution", "organization", "decision making", "critical thinking",
    "collaboration", "negotiation", "multitasking", "initiative", "patience"
}

_CERT_WORDS = {
    "certificate", "certified", "certification", "diploma", "license", "licence",
    "training", "qualification", "course", "degree"
}

_BULLET_SPLIT = re.compile(r"[•·\-\u2013\u2014]+|\n")
_ITEM_SPLIT   = re.compile(r"[,/;•·\|\t]+")
_WORD         = re.compile(r"[A-Za-z][A-Za-z\-\s]+")

def _is_acronym(tok: str) -> bool:
    return tok.isupper() and 2 <= len(tok) <= 6 and tok.isalpha()

def _looks_skill_like(tok: str) -> bool:
    if not tok:
        return False
    t = tok.strip()

    if t.lower() in _STOPWORDS:
        return False
    if re.match(r"^[A-Za-z][A-Za-z\- ]{2,}$", t):
        return True
    if " " in t and not any(ch.isdigit() for ch in t):
        return True
    if t.isupper() and len(t) <= 5:
        return True

    return False

def _normalize_skill(tok: str) -> str:
    t = tok.strip()
    t = re.sub(r"\s+", " ", t)
    t = t.strip(" .,/;|-")
    return t.lower()


def _harvest_from_section(text: str, section_regex: str, max_lines: int = 30) -> list[str]:
    skills = []
    t = text or ""
    m = re.search(section_regex, t, flags=re.I)
    if not m:
        return skills
    after = t[m.end():].splitlines()
    for ln in after[:max_lines]:
        if re.match(r"^\s*[A-Z][A-Za-z &/+-]{0,30}\s*:\s*$", ln):
            break
        for part in _ITEM_SPLIT.split(ln):
            for tok in _WORD.findall(part):
                if _looks_skill_like(tok):
                    skills.append(_normalize_skill(tok))
    return skills


_MONTHS = {
    "jan":1,"feb":2,"mar":3,"apr":4,"may":5,"jun":6,"jul":7,"aug":8,"sep":9,"sept":9,"oct":10,"nov":11,"dec":12
}

def _month_num(s):
    if not s:
        return None
    return _MONTHS.get(s.lower()[:3])

def extract_experience_years_robust(cv_text: str) -> float:
    """Estimate total years of work experience from unstructured text."""
    t = cv_text or ""
    now = datetime.now()
    total_years = 0.0

    numeric_matches = re.findall(
        r"(?i)\b(?:over|more than|at least|approx\.?|around)?\s*(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)\b",
        t
    )
    if numeric_matches:
        numeric_years = max(float(x) for x in numeric_matches)
        total_years = max(total_years, numeric_years)

    since_matches = re.findall(r"(?i)\bsince\s+(?:[A-Za-z]+\s+)?(\d{4})", t)
    for year_str in since_matches:
        try:
            year_val = int(year_str)
            if 1950 < year_val <= now.year:
                total_years = max(total_years, now.year - year_val)
        except Exception:
            pass

    range_pat = re.compile(
        r"(?i)\b(?:(jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)\s+)?(\d{4})\s*[–—\-]\s*"
        r"(?:(jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)\s+)?(present|current|now|\d{4})\b"
    )

    def _mon(m): 
        return {"jan":1,"feb":2,"mar":3,"apr":4,"may":5,"jun":6,"jul":7,
                "aug":8,"sep":9,"sept":9,"oct":10,"nov":11,"dec":12}.get((m or "").lower()[:3])

    for m in range_pat.finditer(t):
        m1, y1s, m2, y2s = m.group(1), m.group(2), m.group(3), m.group(4)
        try:
            y1 = int(y1s)
            y2 = now.year if re.match(r"(?i)present|current|now", y2s) else int(y2s)
            if y2 < y1:
                continue
            start_month = _mon(m1) or 1
            end_month = _mon(m2) or (now.month if re.match(r"(?i)present|current|now", y2s) else 12)
            years = (y2 - y1) + (end_month - start_month) / 12.0
            total_years = max(total_years, years)
        except Exception:
            continue

    month_matches = re.findall(r"(?i)\b(\d+(?:\.\d+)?)\s*(?:months?|mos?)\b", t)
    months_total = sum(float(x) for x in month_matches)
    total_years += months_total / 12.0

    return round(total_years, 1)




NAME_STOPLINES = {
    "curriculum vitae", "resume", "résumé", "profile", "personal profile", "summary", "professional summary",
    "education", "experience", "work experience", "skills", "soft skills", "projects", "contact",
    "contacts", "contact information", "certifications", "interests", "hobbies", "objective"
}

def _clean_token(tok: str) -> str:
    tok = re.sub(r"[^A-Za-z\-\.' ]", "", tok)
    return tok.strip()

def _looks_like_name(line: str) -> bool:
    raw = _clean_token(line)
    if not raw:
        return False
    low = raw.lower()

    if low in NAME_STOPLINES:
        return False
    if any(low.startswith(h) for h in ("address", "phone", "email", "linkedin", "website")):
        return False

    parts = [p for p in raw.split() if p]
    if len(parts) < 2 or len(parts) > 4:
        return False
    if any(re.search(r"\d", p) for p in parts):
        return False

    def good_word(w):
        return re.match(r"^[A-Za-z][A-Za-z\-\.' ]*$", w) is not None

    if not all(good_word(p) for p in parts):
        return False

    joined = " ".join(parts).lower()
    if joined in NAME_STOPLINES:
        return False

    return True

def _normalise_name(s: str) -> str:
    s = " ".join(_clean_token(s).split())
    if s.isupper():
        s = s.title()
    s = " ".join(w[0:1].upper() + w[1:] if w else "" for w in s.split())
    return s.strip()

def _strip_leading_cv_prefix(s: str) -> str:
    """Remove leading 'CV', 'Resume', or 'Curriculum Vitae' prefixes."""
    return re.sub(r"(?i)^\s*(cv|resume|curriculum\s+vitae)[\s_\-:–—]*", "", s).strip()

def _name_from_filename(filename: str) -> str:
    stem = Path(filename).stem
    stem = _strip_leading_cv_prefix(stem)
    stem = re.sub(r"[_\.\-]+", " ", stem)
    stem = " ".join(t for t in stem.split() if not re.fullmatch(r"\d+", t)).strip()
    stem = _strip_leading_cv_prefix(stem)

    parts = [p for p in stem.split() if p]

    for n in range(4, 1, -1):
        for i in range(0, max(0, len(parts) - n + 1)):
            candidate = " ".join(parts[i:i+n])
            if _looks_like_name(candidate):
                return _normalise_name(candidate)

    if len(parts) >= 2:
        return _normalise_name(" ".join(parts[:2]))

    return "N/A"

def _strip_junk_suffixes(s: str) -> str:
    """Remove common CV section headers that get glued to names."""
    for suffix in _JUNK_SUFFIXES:
        s = re.sub(rf"(?i){re.escape(suffix)}.*$", "", s)
        s = re.sub(rf"(?i)([a-z])({re.escape(suffix)})$", r"\1", s)
    return s.strip()


def extract_candidate_name(cv_text: str, filename: str) -> str:
    text = cv_text or ""

    m = re.search(r"(?im)^\s*name\s*[:\-]\s*([A-Za-z][A-Za-z\-\.' ]+(?:\s+[A-Za-z\-\.' ]+){0,3})\s*$", text)
    if m:
        candidate = _strip_junk_suffixes(_strip_leading_cv_prefix(m.group(1)))
        if _looks_like_name(candidate):
            return _normalise_name(candidate)

    lines = text.splitlines()
    for i, ln in enumerate(lines[:40]):
        if re.match(r"(?i)^\s*(curriculum\s+vitae|resume|résumé)\s*$", ln.strip()):
            if i + 1 < len(lines):
                nxt = _strip_junk_suffixes(_strip_leading_cv_prefix(lines[i+1].strip()))
                if _looks_like_name(nxt):
                    return _normalise_name(nxt)

    non_empty = [ln.strip() for ln in lines if ln.strip()]
    for ln in non_empty[:25]:
        candidate = _strip_junk_suffixes(_strip_leading_cv_prefix(ln))
        if _looks_like_name(candidate):
            return _normalise_name(candidate)

    return _name_from_filename(filename)

def fix_candidate_name_in_text(name: str, text: str) -> str:
    """Ensure there is at most ONE 'Candidate Name:' line in the text."""
    text = text or ""
    name = _strip_leading_cv_prefix(name or "")

    pattern = r"(?im)^\s*(?:\*\*)?\s*Candidate\s+Name\s*(?:\*\*)?\s*:\s*.*?$"
    text_clean = re.sub(pattern, "", text, flags=re.MULTILINE)
    text_clean = re.sub(r"\n{3,}", "\n\n", text_clean).strip()

    if name:
        fixed = _normalise_name(name)
        return f"Candidate Name: {fixed}\n\n{text_clean}" if text_clean else f"Candidate Name: {fixed}"
    else:
        return text_clean

_JUNK_SUFFIXES = [
    "professional summary", "professional profile", "professional experience",
    "professional", "summary", "profile", "objective", "experience",
    "curriculum vitae", "resume", "cv", "contact", "about me", "about"
]


def clean_candidate_name_for_table(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"^(?:name[\s:\-]+)+", "", s, flags=re.I)
    
    for suffix in _JUNK_SUFFIXES:
        pattern = rf"(?i){re.escape(suffix)}.*$"
        s = re.sub(pattern, "", s)
        glued_pattern = rf"(?i)([a-z])({re.escape(suffix)})$"
        s = re.sub(glued_pattern, r"\1", s)
    
    s = re.sub(r"[^A-Za-z\-\.' ]+", " ", s)
    s = " ".join(s.split())
    parts = s.split()[:3]
    s = " ".join(parts)
    return _normalise_name(s)


def _parse_match_percent(text: str) -> int:
    if not isinstance(text, str) or not text.strip():
        return 0
    m = re.search(r'(?i)matching\s*(?:percentage|score|match)\s*[:\-]?\s*(\d{1,3})\s*%?', text)
    if m:
        return max(0, min(100, int(m.group(1))))
    nums = [int(n) for n in re.findall(r'\b(100|[1-9]?\d)\b', text)]  # only 0–100
    return max(nums) if nums else 0


def safe_get_match_from_db(filename: str):
    c.execute("SELECT match_explanation FROM CVs WHERE filename = ?", (filename,))
    row = c.fetchone()
    return row[0] if row else None

def _sanitize_collection_name(name: str) -> str:
    s = re.sub(r"[^a-z0-9_]", "_", name.lower())
    if not re.match(r"^[a-z]", s):
        s = "c_" + s
    return s[:63]

def robust_create_or_load_vector_store(filename: str, chunks):
    """Create or load FAISS vector store for a CV."""
    store_path = os.path.join(VECTOR_DIR, filename.replace(".", "_"))
    texts = [doc.page_content for doc in chunks]
    metadatas = [getattr(doc, "metadata", {}) for doc in chunks]

    try:
        if not os.path.exists(store_path):
            vs = FAISS.from_texts(texts=texts, embedding=embedding_model, metadatas=metadatas)
            vs.save_local(store_path)
        else:
            vs = FAISS.load_local(store_path, embeddings=embedding_model, allow_dangerous_deserialization=True)
        return vs
    except Exception:
        shutil.rmtree(store_path, ignore_errors=True)
        vs = FAISS.from_texts(texts=texts, embedding=embedding_model, metadatas=metadatas)
        vs.save_local(store_path)
        return vs



def display_tabs():
    job_desc = st.session_state.job_desc
    CV_texts = st.session_state.CV_texts
    analysis_results = st.session_state.analysis_results

    tab1, tab2 = st.tabs(["Individual Analysis", "Dashboard"])

    with tab1:
        options = list(CV_texts.keys())
        default_ix = 0 if options else None
        selected_file = st.selectbox("Select a CV for analysis:", options, index=default_ix, key="select_cv_file")
        if selected_file and job_desc:
            text = CV_texts[selected_file]
            file_path = os.path.join(UPLOAD_DIR, selected_file)
            chunks = extract_CV_chunks_from_path(file_path)
            vectorstore = robust_create_or_load_vector_store(selected_file, chunks)
            retrieved = retrieve_matching_chunks(job_desc, vectorstore)
            detected_name = extract_candidate_name(text, selected_file)

            response = ""

            if st.button("Analyze CV"):
                with st.spinner("Analysing..."):
                    analysis_text = get_gemini_response(prompt_analysis, retrieved, job_desc)
                    response = fix_candidate_name_in_text(detected_name, analysis_text)




            if response:
                st.write(response)
        else:
            st.info("Please select a CV and paste a job description.")

        with tab2:
            if analysis_results:
                df = pd.DataFrame(analysis_results)

                def safe_parse_match_percent(x):
                    try:
                        if isinstance(x, str):
                            match = re.search(r'\d+', x)
                            if match:
                                return int(match.group(0))
                        return 0
                    except Exception:
                        return 0

                df["Parsed Match %"] = df["Job Match Percentage"].apply(safe_parse_match_percent)

                df["Experience"] = pd.to_numeric(df.get("Experience"), errors="coerce").fillna(0).astype(int)

                df_sorted = df.sort_values(by="Parsed Match %", ascending=False).drop(columns=["Parsed Match %"])
                render_dashboard(df_sorted)
            else:
                st.info("Upload CVs to view dashboard analysis.")


def compute_and_store_match(actual_filename: str, text: str, job_description: str, name_hint: str):
    try:
        chunks = cached_chunks(os.path.join(UPLOAD_DIR, actual_filename))
        vectorstore = cached_faiss(actual_filename, chunks)
        retrieved = cached_retrieve(job_description, vectorstore)
    except Exception:
        retrieved = []

    display_name = clean_candidate_name_for_table(name_hint)

    jd_lower = job_description.lower()
    cv_lower = text.lower()

    jd_words = set(re.findall(r"\b[a-zA-Z]{3,}\b", jd_lower))
    cv_words = set(re.findall(r"\b[a-zA-Z]{3,}\b", cv_lower))
    keyword_overlap = (len(jd_words & cv_words) / max(1, len(jd_words))) * 100

    dynamic_skills = st.session_state.get("jd_skills", [])
    skill_count, _ = count_matching_skills(text, dynamic_skills)
    skill_match = (skill_count / len(dynamic_skills) * 100) if dynamic_skills else 0

    years_est = extract_experience_years_robust(text) or 0
    job_years = 0
    m = re.search(r"(\d+)\s*\+?\s*(?:years?|yrs?)", jd_lower)
    if m:
        job_years = float(m.group(1))
    if job_years > 0:
        experience_match = min((years_est / job_years) * 100, 100)
    else:
        experience_match = min(years_est * 10, 100)

    soft_jd = sum(1 for w in _SOFT_SKILLS if w in jd_lower)
    soft_cv = sum(1 for w in _SOFT_SKILLS if w in cv_lower)
    soft_skill_match = (soft_cv / max(1, soft_jd)) * 100

    try:
        cv_embedding = embedding_model.embed_documents([text[:1500]])[0]
        jd_embedding = cached_embed(job_description)
        dot = sum(a*b for a, b in zip(cv_embedding, jd_embedding))
        norm_a = math.sqrt(sum(a*a for a in cv_embedding))
        norm_b = math.sqrt(sum(b*b for b in jd_embedding))
        embedding_similarity = dot / (norm_a * norm_b + 1e-9)
    except Exception:
        embedding_similarity = 0.0

    embedding_score = max(0, min(100, embedding_similarity * 100))

    Final_Match_Score = (
        0.35 * keyword_overlap +
        0.30 * embedding_score +
        0.15 * skill_match +
        0.15 * experience_match +
        0.05 * soft_skill_match
    )
    
    if Final_Match_Score > 5:
        Final_Match_Score = min(Final_Match_Score * 1.3 + 20, 100)

    match_percent = f"{min(Final_Match_Score, 100):.1f}%"

    match_response = get_gemini_response(prompt_match, retrieved, job_description)
    match_response = fix_candidate_name_in_text(display_name, match_response or "Analysis unavailable.")
    word_count = len((text or "").split())

    c.execute(
        "INSERT OR REPLACE INTO CVs(filename, candidate_name, match_percent, word_count, skills, match_explanation) "
        "VALUES (?, ?, ?, ?, COALESCE((SELECT skills FROM CVs WHERE filename=?), ''), ?)",
        (actual_filename, display_name, match_percent, word_count, actual_filename, match_response)
    )
    conn.commit()
    return match_percent, match_response


def render_login_page():
    """Render the login page."""
    # Center the form with CSS
    st.markdown("""
    <style>
    .main .block-container {
        max-width: 600px;
        padding-top: 2rem;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Create centered container
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.title("🔐 Login")
        st.markdown("---")
        
        with st.form("login_form"):
            username = st.text_input("Username or Email", placeholder="Enter your username or email")
            password = st.text_input("Password", type="password", placeholder="Enter your password")
            submit = st.form_submit_button("Login", use_container_width=True)
            
            if submit:
                if not username or not password:
                    st.error("Please fill in all fields")
                else:
                    success, user_data, message = authenticate_user(username, password)
                    if success:
                        st.session_state.authenticated = True
                        st.session_state.user = user_data
                        st.success(message)
                        st.rerun()
                    else:
                        st.error(message)
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Create Account", use_container_width=True):
                st.session_state.auth_page = "signup"
                st.rerun()
        with col2:
            if st.button("Forgot Password?", use_container_width=True):
                st.session_state.auth_page = "reset_password"
                st.rerun()


def render_signup_page():
    """Render the signup/account creation page."""
    # Center the form with CSS
    st.markdown("""
    <style>
    .main .block-container {
        max-width: 600px;
        padding-top: 2rem;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Create centered container
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.title("📝 Create Account")
        st.markdown("---")
        
        with st.expander("📋 Password Requirements", expanded=False):
            st.markdown("""
            Your password must meet the following requirements:
            - At least 8 characters long
            - At least one uppercase letter (A-Z)
            - At least one lowercase letter (a-z)
            - At least one digit (0-9)
            - At least one special character (!@#$%^&*(),.?":{}|<>)
            """)
        
        with st.form("signup_form"):
            username = st.text_input("Username", placeholder="Choose a username (3-20 characters)")
            email = st.text_input("Email", placeholder="Enter your email address")
            password = st.text_input("Password", type="password", placeholder="Enter a strong password")
            confirm_password = st.text_input("Confirm Password", type="password", placeholder="Confirm your password")
            submit = st.form_submit_button("Create Account", use_container_width=True)
            
            if submit:
                errors = []
                
                # Validate username
                username_valid, username_error = validate_username(username)
                if not username_valid:
                    errors.append(f"Username: {username_error}")
                
                # Validate email
                if not email:
                    errors.append("Email is required")
                elif not validate_email(email):
                    errors.append("Email: Invalid email format")
                
                # Validate password
                if not password:
                    errors.append("Password is required")
                else:
                    password_valid, password_error = validate_password(password)
                    if not password_valid:
                        errors.append(f"Password: {password_error}")
                
                # Check password match
                if password and confirm_password and password != confirm_password:
                    errors.append("Passwords do not match")
                
                if errors:
                    for error in errors:
                        st.error(error)
                else:
                    success, message = create_user(username, email, password)
                    if success:
                        st.success(message)
                        st.info("You can now login with your credentials")
                        st.session_state.auth_page = "login"
                        st.rerun()
                    else:
                        st.error(message)
        
        if st.button("← Back to Login", use_container_width=True):
            st.session_state.auth_page = "login"
            st.rerun()


def render_reset_password_page():
    """Render the password reset page."""
    # Center the form with CSS
    st.markdown("""
    <style>
    .main .block-container {
        max-width: 600px;
        padding-top: 2rem;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Create centered container
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.title("🔑 Reset Password")
        st.markdown("---")
        
        # Check if we have a token in query params or session
        reset_token = st.query_params.get("token") or st.session_state.get("reset_token")
        
        if reset_token:
            # Step 2: Reset password with token
            st.info("Enter your new password below")
            with st.form("reset_password_form"):
                new_password = st.text_input("New Password", type="password", placeholder="Enter a strong password")
                confirm_password = st.text_input("Confirm New Password", type="password", placeholder="Confirm your password")
                submit = st.form_submit_button("Reset Password", use_container_width=True)
                
                if submit:
                    errors = []
                    
                    if not new_password:
                        errors.append("Password is required")
                    else:
                        password_valid, password_error = validate_password(new_password)
                        if not password_valid:
                            errors.append(f"Password: {password_error}")
                    
                    if new_password and confirm_password and new_password != confirm_password:
                        errors.append("Passwords do not match")
                    
                    if errors:
                        for error in errors:
                            st.error(error)
                    else:
                        success, message = reset_password(reset_token, new_password)
                        if success:
                            st.success(message)
                            st.info("You can now login with your new password")
                            st.session_state.auth_page = "login"
                            st.session_state.reset_token = None
                            # Clear query params if present
                            if st.query_params.get("token"):
                                st.query_params.clear()
                            st.rerun()
                        else:
                            st.error(message)
        else:
            # Step 1: Request reset token
            st.info("Enter your email address to receive a password reset link")
            with st.form("request_reset_form"):
                email = st.text_input("Email", placeholder="Enter your email address")
                submit = st.form_submit_button("Send Reset Link", use_container_width=True)
                
                if submit:
                    if not email:
                        st.error("Email is required")
                    elif not validate_email(email):
                        st.error("Invalid email format")
                    else:
                        success, token, message = generate_reset_token(email)
                        if success:
                            # In a real app, you would send this token via email
                            # For now, we'll display it (in production, remove this!)
                            st.success("Password reset token generated!")
                            st.warning("⚠️ In production, this token would be sent via email. For now, copy this token:")
                            st.code(token, language=None)
                            st.info("Add ?token=YOUR_TOKEN to the URL or use the token below")
                            
                            # Store token in session for convenience
                            st.session_state.reset_token = token
                            st.rerun()
                        else:
                            st.error(message)
        
        if st.button("← Back to Login", use_container_width=True):
            st.session_state.auth_page = "login"
            st.session_state.reset_token = None
            st.rerun()


# Authentication check - show login if not authenticated
if not st.session_state.authenticated:
    # Show appropriate auth page
    if st.session_state.auth_page == "signup":
        render_signup_page()
    elif st.session_state.auth_page == "reset_password":
        render_reset_password_page()
    else:
        render_login_page()
    
    st.stop()  # Stop execution here if not authenticated

# Main app content (only shown if authenticated)
st.title("CV Analysis System")
st.subheader("Curriculum Vitae Parsing and Analysis System Using Natural Language Processing and Document AI")

# Logout button in sidebar
with st.sidebar:
    st.markdown("---")
    st.write(f"**Logged in as:** {st.session_state.user['username']}")
    if st.button("Logout", use_container_width=True):
        st.session_state.authenticated = False
        st.session_state.user = None
        st.rerun()

if "jd_key" not in st.session_state: st.session_state.jd_key = 0
if "uploader_key" not in st.session_state: st.session_state.uploader_key = 0

job_desc = st.sidebar.text_area(
    "Add Job Description:",
    height=150,
    key=f"jd_{st.session_state.jd_key}"
)
uploaded_files = st.sidebar.file_uploader(
    "Upload CVs",
    type=['pdf', 'docx', 'txt'],
    accept_multiple_files=True,
    key=f"uploader_{st.session_state.uploader_key}"
)


col1, col2 = st.sidebar.columns(2)
with col1:
    submit_clicked = st.button("Submit", use_container_width=True)
with col2:
    if st.button("Clear", use_container_width=True):
        clear_cache(flag=0)

if submit_clicked:
    dynamic_skills = extract_required_skills(job_desc) if job_desc else []
    st.session_state.jd_skills = dynamic_skills

    
    if not job_desc or not uploaded_files:
        st.warning("Please provide job description and upload at least one CV.")
    else:
        time.sleep(1)
        st.session_state.submitted = True
        st.session_state.job_desc = job_desc
        st.session_state.uploaded_files = uploaded_files
        st.session_state.CV_texts = {}
        st.session_state.analysis_results = []

        progress_bar = st.progress(0, text="Processing CVs...")
        total_files = len(uploaded_files)

        used_names = set()
        for idx, file in enumerate(uploaded_files):
            progress_bar.progress((idx + 1) / total_files, text=f"Processing {file.name} ({idx + 1}/{total_files})")
            saved_path = save_uploaded_file(file)
            actual_filename = os.path.basename(saved_path)
            if actual_filename in used_names:
                continue
            used_names.add(actual_filename)

            file_bytes = Path(saved_path).read_bytes()
            file_type = getattr(file, "type", None)
            if not file_type:
                ext = os.path.splitext(file.name)[1].lower()
                if ext == ".pdf":
                    file_type = "application/pdf"
                elif ext == ".docx":
                    file_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                elif ext == ".txt":
                    file_type = "text/plain"
                else:
                    file_type = "application/octet-stream"

            text = cached_extract_text(file_bytes, file_type)
            st.session_state.CV_texts[actual_filename] = text

            skill_count, skill_list = count_matching_skills(text, dynamic_skills)
            cv_skills = extract_skills_from_cv(text)

            detected_name = extract_candidate_name(text, actual_filename)

            c.execute("SELECT filename, candidate_name, match_percent, word_count, skills, match_explanation FROM CVs WHERE filename = ?", (actual_filename,))
            cached = c.fetchone()

            word_count = len(text.split())

            if cached:
                _, cached_name, cached_pct, _, _, cached_expl = cached
                name_for_row = cached_name if (cached_name and cached_name != "N/A") else detected_name

                # Recompute only if no cached % or it parses to 0
                need_recompute = (not cached_pct) or (_parse_match_percent(str(cached_pct)) == 0)

                if need_recompute:
                    match_percent, match_response = compute_and_store_match(actual_filename, text, job_desc, name_for_row)
                else:
                    match_percent = cached_pct
                    match_response = fix_candidate_name_in_text(name_for_row, cached_expl or "")
                    try:
                        c.execute("UPDATE CVs SET candidate_name=?, match_explanation=? WHERE filename=?",
                                (clean_candidate_name_for_table(name_for_row), match_response, actual_filename))
                        conn.commit()
                    except Exception:
                        pass

            else:
                name_for_row = detected_name
                match_percent, match_response = compute_and_store_match(actual_filename, text, job_desc, name_for_row)

            jd_lower = job_desc.lower()
            cv_lower = text.lower()

            jd_keywords = set(re.findall(r"\b[a-zA-Z]{3,}\b", jd_lower))
            cv_keywords = set(re.findall(r"\b[a-zA-Z]{3,}\b", cv_lower))
            keyword_overlap = (len(jd_keywords & cv_keywords) / len(jd_keywords)) * 100 if jd_keywords else 0

            skill_match = (skill_count / len(dynamic_skills)) * 100 if dynamic_skills else 0

            years_est = extract_experience_years_robust(text)
            exp_in_jd = re.findall(r"(\d+)\s*(?:\+?\s*)?(?:years?|yrs?)", jd_lower)
            if exp_in_jd:
                jd_exp = max(float(x) for x in exp_in_jd)
                experience_match = min((years_est / jd_exp) * 100, 100)
            else:
                experience_match = min(years_est * 10, 100)

            cert_jd = sum(1 for w in _CERT_WORDS if w in jd_lower)
            cert_cv = sum(1 for w in _CERT_WORDS if w in cv_lower)
            certificate_match = (cert_cv / cert_jd) * 100 if cert_jd > 0 else 0

            soft_jd = sum(1 for w in _SOFT_SKILLS if w in jd_lower)
            soft_cv = sum(1 for w in _SOFT_SKILLS if w in cv_lower)
            soft_skill_match = (soft_cv / soft_jd) * 100 if soft_jd > 0 else 0

            st.session_state.analysis_results.append({
                "Filename": actual_filename,
                "Candidate Name": clean_candidate_name_for_table(name_for_row).upper(),
                "Job Match Percentage": match_percent if cached else (match_percent if 'match_percent' in locals() else "0%"),
                "Detected Skills": cv_skills if cv_skills else [],
                "Experience": round(years_est, 1)
            })

            eval_row = {
                "Timestamp": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "Filename": file.name,
                "Candidate Name": name_for_row,
                "Match %": match_percent if cached else (match_percent if 'match_percent' in locals() else "0%"),
                "Keyword Overlap %": round(keyword_overlap, 1),
                "Skill Match %": round(skill_match, 1),
                "Experience Match %": round(experience_match, 1),
                "Certificate Match %": round(certificate_match, 1),
                "Soft Skill Match %": round(soft_skill_match, 1)
            }

            eval_df = pd.DataFrame([eval_row])
            eval_file = "evaluation_results.csv"
            if os.path.exists(eval_file):
                eval_df.to_csv(eval_file, mode="a", index=False, header=False)
            else:
                eval_df.to_csv(eval_file, index=False)


        progress_bar.empty()
        if st.session_state.submitted and st.session_state.analysis_results:
            display_tabs()
elif st.session_state.submitted and st.session_state.analysis_results:
    display_tabs()
else:
    st.info("Please select a CV and paste a job description.")


def _parse_pdf_date(date_str: str) -> datetime | None:
    """Parse PDF date format like D:20180628083536+02'00' to datetime."""
    if not date_str:
        return None
    try:
        clean = date_str.replace("D:", "").replace("'", "")
        if "+" in clean:
            clean = clean.split("+")[0]
        elif "-" in clean and len(clean) > 14:
            clean = clean[:14]
        return datetime.strptime(clean[:14], "%Y%m%d%H%M%S")
    except Exception:
        return None


def verify_document_integrity(file_bytes: bytes, filename: str) -> dict:
    """
    Certificate verification focusing on TAMPERING indicators, not creation tools.
    
    Key principle: Creating a certificate with design software is NORMAL.
    What's suspicious is MODIFYING an existing certificate later.
    """
    result = {"type": "unknown", "risk": "PASS", "flags": {}, "notes": [], "reasons": []}
    suspicious_score = 0

    kind = filetype.guess(file_bytes)
    if kind is None:
        result["risk"] = "UNKNOWN"
        result["reasons"].append("Could not detect file type")
        return result

    result["type"] = kind.mime

    if "image" in kind.mime:
        result["notes"].append(f"Format: {kind.mime.upper()}")
        result["notes"].append(f"File size: {len(file_bytes) / 1024:.1f} KB")
        
        img_data = file_bytes.decode('latin-1', errors='ignore')
        
        has_xmp = b"<x:xmpmeta" in file_bytes or b"xmp" in file_bytes.lower()
        if has_xmp and "ModifyDate" in img_data:
            create_match = re.search(r'CreateDate["\'>:]+(\d{4})', img_data)
            modify_match = re.search(r'ModifyDate["\'>:]+(\d{4})', img_data)
            
            if create_match and modify_match:
                create_year = int(create_match.group(1))
                modify_year = int(modify_match.group(1))
                
                if modify_year > create_year:
                    suspicious_score += 3
                    result["flags"]["Modified After Creation"] = f"Created {create_year}, modified {modify_year}"
                    result["reasons"].append(f"Image was modified in {modify_year}, years after creation in {create_year}")
        
        if len(file_bytes) < 5000:
            suspicious_score += 1
            result["reasons"].append("Very small file size - may indicate heavy compression or editing")
        
        if suspicious_score == 0:
            result["reasons"].append("Image metadata appears consistent")
            result["reasons"].append("No signs of post-creation tampering detected")

    elif "pdf" in kind.mime:
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            meta = doc.metadata or {}
            page_count = doc.page_count
            doc.close()

            producer = meta.get("producer", "") or ""
            creator = meta.get("creator", "") or ""
            
            result["notes"].append(f"Pages: {page_count}")
            if meta.get("author"):
                result["notes"].append(f"Author: {meta.get('author')}")
            result["notes"].append(f"Creator: {creator or 'N/A'}")
            result["notes"].append(f"Producer: {producer or 'N/A'}")
            result["notes"].append(f"Created: {meta.get('creationDate', 'N/A')}")
            result["notes"].append(f"Modified: {meta.get('modDate', 'N/A')}")

            creation_date = _parse_pdf_date(meta.get("creationDate", ""))
            mod_date = _parse_pdf_date(meta.get("modDate", ""))
            
            if creation_date and mod_date:
                time_diff = (mod_date - creation_date).days
                years_diff = time_diff // 365
                
                if years_diff >= 1:
                    suspicious_score += 3
                    result["flags"]["Significant Date Gap"] = f"Modified {years_diff}+ years after creation"
                    result["reasons"].append(f"Document modified {years_diff} year(s) after original creation - likely tampered")
                elif time_diff > 90:
                    suspicious_score += 2
                    result["flags"]["Date Gap"] = f"Modified {time_diff} days after creation"
                    result["reasons"].append(f"Document modified {time_diff} days after creation - review recommended")
                elif time_diff <= 1:
                    result["reasons"].append("Creation and modification dates are consistent (same day)")
            
            if creation_date and not mod_date:
                result["reasons"].append("Document has not been modified since creation")
            
            online_editors = ["smallpdf", "ilovepdf", "sejda", "pdf24", "pdfcandy", "sodapdf"]
            producer_lower = producer.lower()
            creator_lower = creator.lower()
            
            for editor in online_editors:
                if editor in producer_lower or editor in creator_lower:
                    suspicious_score += 2
                    result["flags"]["Online PDF Editor"] = f"Processed with: {editor}"
                    result["reasons"].append(f"Document was processed through online PDF editor ({editor})")
                    break
            
            if not meta.get("creationDate") and not meta.get("modDate") and not meta.get("producer"):
                suspicious_score += 2
                result["flags"]["Stripped Metadata"] = "Document metadata appears to be removed"
                result["reasons"].append("Document metadata has been stripped - common in tampered documents")
            
            if suspicious_score == 0:
                result["reasons"].append("Document metadata is consistent and intact")
                result["reasons"].append("No indicators of post-creation tampering")
                if "adobe" in producer_lower or "adobe" in creator_lower:
                    result["reasons"].append("Created with Adobe software (professional standard)")
                elif "microsoft" in producer_lower or "word" in creator_lower:
                    result["reasons"].append("Created with Microsoft Office (common for certificates)")
                elif "photoshop" in producer_lower or "photoshop" in creator_lower:
                    result["reasons"].append("Designed in Photoshop (professional design tool)")
                elif "illustrator" in producer_lower or "illustrator" in creator_lower:
                    result["reasons"].append("Designed in Illustrator (professional design tool)")

        except Exception as e:
            result["risk"] = "ERROR"
            result["reasons"].append(f"Could not analyze PDF: {str(e)}")
            return result
    else:
        result["risk"] = "UNSUPPORTED"
        result["reasons"].append(f"File type '{kind.mime}' is not supported")
        return result

    if suspicious_score >= 3:
        result["risk"] = "SUSPICIOUS"
    elif suspicious_score >= 2:
        result["risk"] = "REVIEW"
    else:
        result["risk"] = "PASS"

    return result


with st.expander("Certificate / Document Verification", expanded=False):
    st.caption("Upload a certificate (PDF or image) to check for signs of tampering or editing.")
    cert_file = st.file_uploader("Upload certificate to verify", type=["pdf", "jpg", "jpeg", "png"], key="verify_uploader")
    
    if cert_file is not None:
        file_bytes = cert_file.getvalue()
        report = verify_document_integrity(file_bytes, cert_file.name)

        st.markdown("---")
        st.write(f"**File:** {cert_file.name}")
        st.write(f"**Type:** {report['type'].upper()}")
        
        if report["risk"] == "PASS":
            st.success("Verdict: PASS")
        elif report["risk"] == "REVIEW":
            st.warning("Verdict: NEEDS REVIEW")
        elif report["risk"] == "SUSPICIOUS":
            st.error("Verdict: SUSPICIOUS")
        else:
            st.warning(f"Verdict: {report['risk']}")

        if report["reasons"]:
            st.subheader("Analysis")
            for reason in report["reasons"]:
                st.write(f"- {reason}")

        if report["flags"]:
            st.subheader("Flags Detected")
            for k, v in report["flags"].items():
                st.write(f"- **{k}**: {v}")

        with st.expander("Technical Details", expanded=False):
            for note in report["notes"]:
                st.write(f"- {note}")